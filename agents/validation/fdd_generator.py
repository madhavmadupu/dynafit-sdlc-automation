"""
agents/validation/fdd_generator.py
Phase 5 FDD (Functional Design Document) Generator.
Produces a downloadable .docx document from validated fitment results.
"""

from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

import structlog
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from core.config.settings import settings
from core.schemas.classification_result import ClassificationResult, ValidatedFitmentBatch
from core.schemas.enums import Verdict
from core.schemas.requirement_atom import RequirementAtom

log = structlog.get_logger()

# ── Color palette ───────────────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)  # Brand blue
COLOR_FIT = RGBColor(0x05, 0x96, 0x69)  # Emerald
COLOR_PARTIAL = RGBColor(0xD9, 0x73, 0x06)  # Amber
COLOR_GAP = RGBColor(0xDC, 0x26, 0x26)  # Red
COLOR_GRAY = RGBColor(0x64, 0x74, 0x8B)  # Slate-500
COLOR_DARK = RGBColor(0x1E, 0x29, 0x3B)  # Slate-800
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

VERDICT_COLORS = {
    Verdict.FIT: COLOR_FIT,
    Verdict.PARTIAL_FIT: COLOR_PARTIAL,
    Verdict.GAP: COLOR_GAP,
}

VERDICT_LABELS = {
    Verdict.FIT: "FIT — Standard D365 Feature",
    Verdict.PARTIAL_FIT: "PARTIAL FIT — Configuration Required",
    Verdict.GAP: "GAP — Custom Development Required",
}


def _set_cell_shading(cell, color_hex: str) -> None:  # noqa: ANN001
    """Apply background shading to a table cell."""
    shading = cell._tc.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        },
    )
    shading.append(shading_elem)


def _add_styled_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 11,
    color: RGBColor = COLOR_DARK,
    space_after: int = 6,
    alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"


def _build_summary_table(
    doc: Document,
    fit_count: int,
    partial_count: int,
    gap_count: int,
    total: int,
    override_count: int,
) -> None:
    """Add executive summary statistics table."""
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Total Requirements", "FIT", "PARTIAL FIT", "GAP", "Overrides"]
    values = [str(total), str(fit_count), str(partial_count), str(gap_count), str(override_count)]
    header_colors = ["4F81BD", "059669", "D97306", "DC2626", "6366F1"]

    for i, (header, value, bg_color) in enumerate(zip(headers, values, header_colors)):
        # Header row
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_WHITE
        run.font.name = "Calibri"
        _set_cell_shading(cell, bg_color)

        # Value row
        cell = table.cell(1, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_DARK
        run.font.name = "Calibri"

        # Percentage
        if total > 0 and header in ("FIT", "PARTIAL FIT", "GAP"):
            pct = int(value) / total * 100
            pct_run = p.add_run(f"\n({pct:.0f}%)")
            pct_run.font.size = Pt(8)
            pct_run.font.color.rgb = COLOR_GRAY
            pct_run.font.name = "Calibri"


def _add_requirement_entry(
    doc: Document,
    result: ClassificationResult,
    atom: RequirementAtom | None,
    index: int,
) -> None:
    """Add a single requirement entry with its details."""
    verdict_color = VERDICT_COLORS.get(result.verdict, COLOR_DARK)

    # Requirement header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    atom_id_short = str(result.atom_id)[:8]
    run = p.add_run(f"{index}. ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    run.font.name = "Calibri"

    run = p.add_run(f"[{atom_id_short}] ")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_GRAY
    run.font.name = "Calibri"

    # Verdict badge inline
    run = p.add_run(f"  {result.verdict.value}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = verdict_color
    run.font.name = "Calibri"

    # Module & priority
    if atom:
        meta_parts = []
        meta_parts.append(f"Module: {atom.module.value}")
        meta_parts.append(f"Priority: {atom.priority.value}")
        if atom.country:
            meta_parts.append(f"Country: {atom.country}")
        _add_styled_paragraph(doc, " | ".join(meta_parts), size=9, color=COLOR_GRAY, space_after=2)

    # Requirement text
    req_text = atom.text if atom else "(Requirement text not available)"
    _add_styled_paragraph(doc, req_text, size=10, space_after=4)

    # Details table
    details = []
    if result.matched_capability:
        details.append(("D365 Capability", result.matched_capability))
    if result.config_needed:
        details.append(("Configuration Required", result.config_needed))
    if result.gap_description:
        details.append(("Gap Description", result.gap_description))
    if result.caveats:
        details.append(("Caveats", "; ".join(result.caveats)))
    details.append(("Confidence", f"{result.confidence:.0%}"))
    details.append(("AI Rationale", result.rationale))

    if details:
        table = doc.add_table(rows=len(details), cols=2)
        table.style = "Table Grid"
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(14)

        for row_idx, (label, value) in enumerate(details):
            label_cell = table.cell(row_idx, 0)
            label_cell.text = ""
            p = label_cell.paragraphs[0]
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_DARK
            run.font.name = "Calibri"
            _set_cell_shading(label_cell, "F1F5F9")

            value_cell = table.cell(row_idx, 1)
            value_cell.text = ""
            p = value_cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_DARK
            run.font.name = "Calibri"

    doc.add_paragraph()  # spacing


def generate_fdd(
    batch: ValidatedFitmentBatch,
    atoms: list[RequirementAtom],
) -> str:
    """
    Generate the Functional Design Document (.docx).

    Produces a structured document organized by verdict category:
    - Section 1: FIT requirements — standard D365 capabilities
    - Section 2: PARTIAL FIT requirements — config/customization needed
    - Section 3: GAP requirements — custom X++ development needed

    Args:
        batch: Final validation batch containing results and overrides.
        atoms: Original requirement atoms.

    Returns:
        Absolute path to the generated .docx file.
    """
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # ── Default font ──────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Cover / Title ─────────────────────────────────────────────────────
    doc.add_paragraph()  # top spacer

    _add_styled_paragraph(
        doc,
        "FUNCTIONAL DESIGN DOCUMENT",
        bold=True,
        size=24,
        color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    _add_styled_paragraph(
        doc,
        "D365 Finance & Operations — Requirement Fitment Analysis",
        size=13,
        color=COLOR_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=20,
    )

    # Metadata
    stamp = datetime.utcnow().strftime("%B %d, %Y at %H:%M UTC")
    _add_styled_paragraph(
        doc,
        f"Run ID: {batch.run_id}",
        size=9,
        color=COLOR_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    _add_styled_paragraph(
        doc,
        f"Generated: {stamp}",
        size=9,
        color=COLOR_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    _add_styled_paragraph(
        doc,
        f"Status: {batch.run_status.value}",
        size=9,
        color=COLOR_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=20,
    )

    # Horizontal rule
    doc.add_paragraph("_" * 80)

    # ── Executive Summary ─────────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)

    atom_map = {str(a.id): a for a in atoms}
    total = len(batch.results)

    fit_results = [r for r in batch.results if r.verdict == Verdict.FIT]
    partial_results = [r for r in batch.results if r.verdict == Verdict.PARTIAL_FIT]
    gap_results = [r for r in batch.results if r.verdict == Verdict.GAP]

    _add_styled_paragraph(
        doc,
        f"This document presents the fitment analysis of {total} requirement atoms "
        f"against Microsoft Dynamics 365 Finance & Operations capabilities.",
        size=10,
        space_after=8,
    )

    _build_summary_table(
        doc,
        fit_count=len(fit_results),
        partial_count=len(partial_results),
        gap_count=len(gap_results),
        total=total,
        override_count=batch.override_count,
    )

    doc.add_paragraph()  # spacing

    if total > 0:
        fit_rate = len(fit_results) / total * 100
        _add_styled_paragraph(
            doc,
            f"Overall Fit Rate: {fit_rate:.1f}% of requirements can be addressed with "
            f"standard or configurable D365 capabilities. "
            f"{len(gap_results)} requirement(s) require custom X++ development.",
            size=10,
            space_after=12,
        )

    # ── Conflict Summary ──────────────────────────────────────────────────
    if batch.conflict_report and batch.conflict_report.conflicts:
        doc.add_heading("1.1 Conflict Summary", level=2)
        _add_styled_paragraph(
            doc,
            f"{batch.conflict_report.error_count} blocking conflict(s) and "
            f"{batch.conflict_report.warning_count} warning(s) detected.",
            size=10,
            space_after=4,
        )
        for conflict in batch.conflict_report.conflicts:
            _add_styled_paragraph(
                doc,
                f"[{conflict.severity}] {conflict.conflict_type}: {conflict.description}",
                size=9,
                color=COLOR_GAP if conflict.severity == "ERROR" else COLOR_PARTIAL,
                space_after=2,
            )
        doc.add_paragraph()

    # ── Section 2: FIT Requirements ───────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("2. FIT — Standard D365 Features", level=1)
    _add_styled_paragraph(
        doc,
        f"{len(fit_results)} requirement(s) are fully covered by out-of-the-box "
        f"D365 Finance & Operations capabilities. No custom development is needed.",
        size=10,
        space_after=8,
        color=COLOR_FIT,
    )

    if fit_results:
        for i, result in enumerate(fit_results, 1):
            atom = atom_map.get(str(result.atom_id))
            _add_requirement_entry(doc, result, atom, i)
    else:
        _add_styled_paragraph(doc, "No requirements classified as FIT.", size=10, color=COLOR_GRAY)

    # ── Section 3: PARTIAL FIT Requirements ───────────────────────────────
    doc.add_page_break()
    doc.add_heading("3. PARTIAL FIT — Configuration & Customization", level=1)
    _add_styled_paragraph(
        doc,
        f"{len(partial_results)} requirement(s) can be addressed through D365 configuration, "
        f"parameter setup, or minor customization. These do not require full custom development "
        f"but may need functional consultant effort.",
        size=10,
        space_after=8,
        color=COLOR_PARTIAL,
    )

    if partial_results:
        for i, result in enumerate(partial_results, 1):
            atom = atom_map.get(str(result.atom_id))
            _add_requirement_entry(doc, result, atom, i)
    else:
        _add_styled_paragraph(
            doc, "No requirements classified as PARTIAL FIT.", size=10, color=COLOR_GRAY
        )

    # ── Section 4: GAP Requirements ───────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. GAP — Custom X++ Development Required", level=1)
    _add_styled_paragraph(
        doc,
        f"{len(gap_results)} requirement(s) cannot be fulfilled by standard D365 capabilities "
        f"and require custom X++ development, ISV solutions, or third-party integrations.",
        size=10,
        space_after=8,
        color=COLOR_GAP,
    )

    if gap_results:
        for i, result in enumerate(gap_results, 1):
            atom = atom_map.get(str(result.atom_id))
            _add_requirement_entry(doc, result, atom, i)
    else:
        _add_styled_paragraph(doc, "No requirements classified as GAP.", size=10, color=COLOR_GRAY)

    # ── Section 5: Overrides & Audit Trail ────────────────────────────────
    if batch.overrides:
        doc.add_page_break()
        doc.add_heading("5. Consultant Overrides", level=1)
        _add_styled_paragraph(
            doc,
            f"{len(batch.overrides)} classification(s) were overridden by consultant review.",
            size=10,
            space_after=8,
        )

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Atom ID", "Original", "Override", "Reason", "Reviewed By"]
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_WHITE
            run.font.name = "Calibri"
            _set_cell_shading(cell, "4F81BD")

        for override in batch.overrides:
            row = table.add_row()
            values = [
                str(override.atom_id)[:8],
                override.original_verdict.value,
                override.override_verdict.value,
                override.reason,
                override.reviewed_by,
            ]
            for i, value in enumerate(values):
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(value)
                run.font.size = Pt(8)
                run.font.color.rgb = COLOR_DARK
                run.font.name = "Calibri"

    # ── Footer note ───────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph("_" * 80)
    _add_styled_paragraph(
        doc,
        "This document was auto-generated by the DYNAFIT SDLC Automation Engine. "
        "All classifications have been validated through the multi-agent pipeline "
        "and consultant review process.",
        size=8,
        color=COLOR_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── Save ──────────────────────────────────────────────────────────────
    os.makedirs(settings.OUTPUT_DIR, exist_ok=True)
    stamp_file = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filepath = Path(settings.OUTPUT_DIR) / f"fdd_{batch.run_id}_{stamp_file}.docx"

    doc.save(str(filepath))
    log.info("fdd_generated", run_id=batch.run_id, path=str(filepath))

    return str(filepath)
